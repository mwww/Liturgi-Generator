from docx import Document
from .helpers import GetLyricsWeb, GetLyricsFile, GetVerse
import os
base_dir = 'D:\\Data Hansen\\Data Programming\\Projects\\Liturgi Generator\\src\\'

def NewDocument(data):
    # Get User Input
    date = data["date"]
    theme = data["theme"]
    pdt = data["pdt"]
    pnt = data["pnt"]
    verseFirman = data["verseFirman"]
    verseKataPembuka = data["verseKataPembuka"]
    verseBeritaAnugerah = data["verseBeritaAnugerah"]
    versePersembahan = data["versePersembahan"]
    pelayananPujian = data["pelayananPujian"]
        # Make a list of song dicts from the user input
    songList = []
    for songLabels, songBook, songBook_number, songBook_verse, songTitle, songLyrics in data['songZip']:
        songDict = {
            'songBook': songBook,
            'songBook_number': songBook_number,
            'songBook_verse': songBook_verse,
            'songTitle': songTitle,
            'songLyrics': songLyrics.replace('\r', ''), #Modify song lyrics format from \r\n to \n only, if exist
            'completeTitle': ''
        }
        if songBook:
            # songDict['completeTitle'] = f'{songDict['songBook']} {songDict['songBook_number']} : {songDict['songBook_verse']} "{songDict['songTitle']}"'
            songDict['completeTitle'] = songDict['songBook'] + " " + songDict['songBook_number'] + " : " + songDict['songBook_verse'] + ' "' + songDict['songTitle'] + '"'
            songDict['songLyrics'] = GetLyricsFile(songDict) #Search for lyrics in the book, if book exist
        else:
            songDict['completeTitle'] = songDict['songTitle']
        songList.append(songDict)


    # Read Template, Modify, and Save to New Docs
    os.chdir(base_dir)
    document = Document('Liturgi_Template.docx')
    for paragraph in document.paragraphs:
        content = paragraph.text

        # Modify the date
        if content == '[Date]':
            paragraph.text = f'Minggu, {date}'

        # Modify the theme
        if content == '[Theme]':
            paragraph.text = f'"{theme}"'

        # Modify the people
        if '[Pendeta]' in content:
            for run in paragraph.runs:
                if run.text == '[Pendeta]':
                    run.text = f'{pdt}'
        if content == '[Penatua]':
            paragraph.text = f'{pnt}'
        
        # Modify the songs
        if 'Song' in content:
            for i in range(6):
                if f'Song{i+1}' in content:
                    completeTitle = songList[i]['completeTitle']
                    # songLyrics = GetLyricsFile(songList[i])
                    if content == f'[Song{i+1}]':
                        paragraph.text = completeTitle
                    if content == f'[Song{i+1}_Lyrics]':
                        # paragraph.text = songLyrics
                        paragraph.text = songList[i]['songLyrics']
                    break

        # Modify the verses
        if '[Verse_Kata_Pembuka]' in content:
            verseKataPembukaText = GetVerse(verseKataPembuka)
            for run in paragraph.runs:
                if run.text == '[Verse_Kata_Pembuka]':
                    run.text = verseKataPembuka
                if run.text == '[Verse_Kata_Pembuka_Text]':
                    run.text = verseKataPembukaText

        if 'Verse_Berita_Anugerah' in content:
            verseBeritaAnugerahText = GetVerse(verseBeritaAnugerah)
            for run in paragraph.runs:
                if run.text == '[Verse_Berita_Anugerah]':
                    run.text = verseBeritaAnugerah
                if run.text == '[Verse_Berita_Anugerah_Text]':
                    run.text = verseBeritaAnugerahText
        
        if '[Verse_Firman]' in content:
            for run in paragraph.runs:
                if run.text == '[Verse_Firman]':
                    run.text = verseFirman

        if 'Verse_Persembahan' in content:
            versePersembahanText = GetVerse(versePersembahan)
            for run in paragraph.runs:
                if run.text == '[Verse_Persembahan]':
                    run.text = versePersembahan
                if run.text == '[Verse_Persembahan_Text]':
                    run.text = versePersembahanText

        # Add/Remove Pelayanan Pujian
        if '[Pelayanan_Pujian]' in content:
            for run in paragraph.runs:
                if run.text == '[Pelayanan_Pujian]':
                    if pelayananPujian:
                        run.text = '\nPELAYANAN PUJIAN'
                    else:
                        run.text = ''
    
    os.chdir(base_dir)
    document.save('New_Liturgi.docx')